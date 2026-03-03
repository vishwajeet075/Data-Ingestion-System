# document_processor.py
import os
import tempfile
from typing import Tuple, Dict, Any
from pathlib import Path
import streamlit as st
import PyPDF2
import pdfplumber
from pdfminer.high_level import extract_text as pdfminer_extract
from pdfminer.layout import LAParams
import docx
from multimedia_processor import MultimediaProcessor

class DocumentProcessor:
    """Optimized document processing with better text extraction"""

    # Increased limits for better extraction
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB instead of 10MB
    CHUNK_SIZE = 2 * 1024 * 1024  # 2MB chunks for reading

    @staticmethod
    def _read_file_in_chunks(file_path: str, chunk_size: int = 8192) -> bytes:
        """Read file in chunks to handle large files"""
        content = b""
        with open(file_path, 'rb') as f:
            while True:
                chunk = f.read(chunk_size)
                if not chunk:
                    break
                content += chunk
        return content

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using multiple methods - COMPLETE EXTRACTION"""
        debug_info = {
            "methods_tried": [],
            "methods_success": {},
            "total_pages": 0,
            "total_words": 0,
            "total_chars": 0,
            "pages_processed": 0,
            "pages_skipped": 0
        }

        all_text = ""

        # Method 1: pdfplumber (most reliable)
        try:
            with pdfplumber.open(file_path) as pdf:
                debug_info["total_pages"] = len(pdf.pages)
                pages_text = []

                for i, page in enumerate(pdf.pages):
                    page_text = ""

                    # Strategy 1: Standard extraction
                    page_text = page.extract_text()

                    # Strategy 2: Extract with x_tolerance for better spacing
                    if not page_text or len(page_text.strip()) < 5:
                        page_text = page.extract_text(x_tolerance=3, y_tolerance=3)

                    # Strategy 3: Extract by words and join
                    if not page_text or len(page_text.strip()) < 5:
                        words = page.extract_words()
                        if words:
                            page_text = " ".join([word['text'] for word in words])

                    # FIXED: Accept even minimal text instead of skipping
                    if page_text and page_text.strip():  # Changed from > 5 to just check if exists
                        pages_text.append(page_text)
                        debug_info[f"page_{i+1}_words"] = len(page_text.split())
                        debug_info[f"page_{i+1}_chars"] = len(page_text)
                        debug_info["pages_processed"] += 1
                    else:
                        # Mark page as skipped but add placeholder to maintain structure
                        pages_text.append(f"[Page {i+1} - No extractable text]")
                        debug_info["pages_skipped"] += 1
                        st.warning(f"⚠️ Page {i+1} has no extractable text")

                if pages_text:
                    all_text = "\n\n".join(pages_text)  # Double newline for page separation
                    debug_info["methods_tried"].append("pdfplumber")
                    debug_info["methods_success"]["pdfplumber"] = True
                    debug_info["total_words"] = len(all_text.split())
                    debug_info["total_chars"] = len(all_text)
                    st.success(f"✓ pdfplumber: Extracted {debug_info['total_words']} words from {debug_info['pages_processed']}/{debug_info['total_pages']} pages")
                    return all_text.strip(), debug_info

        except Exception as e:
            debug_info["methods_success"]["pdfplumber"] = False
            debug_info["pdfplumber_error"] = str(e)
            st.warning(f"pdfplumber failed: {str(e)[:100]}")

        # Method 2: PyPDF2 (fallback)
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                debug_info["total_pages"] = len(pdf_reader.pages)
                debug_info["pages_processed"] = 0

                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():  # Accept any non-empty text
                            text += f"\n\n--- Page {i+1} ---\n\n{page_text}"
                            debug_info["pages_processed"] += 1
                    except Exception as page_error:
                        st.warning(f"Error on page {i+1}: {str(page_error)[:50]}")
                        continue

                if text.strip():
                    all_text = text
                    debug_info["methods_tried"].append("PyPDF2")
                    debug_info["methods_success"]["PyPDF2"] = True
                    debug_info["total_words"] = len(all_text.split())
                    debug_info["total_chars"] = len(all_text)
                    st.success(f"✓ PyPDF2: Extracted {debug_info['total_words']} words from {debug_info['pages_processed']} pages")
                    return all_text.strip(), debug_info

        except Exception as e:
            debug_info["methods_success"]["PyPDF2"] = False
            debug_info["pypdf2_error"] = str(e)
            st.warning(f"PyPDF2 failed: {str(e)[:100]}")

        # Method 3: pdfminer (for complex layouts)
        try:
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                boxes_flow=0.5,
                detect_vertical=True,  # Better vertical text detection
                all_texts=True  # Extract all text elements
            )

            text = pdfminer_extract(file_path, laparams=laparams)
            if text and text.strip():
                all_text = text
                debug_info["methods_tried"].append("pdfminer")
                debug_info["methods_success"]["pdfminer"] = True
                debug_info["total_words"] = len(all_text.split())
                debug_info["total_chars"] = len(all_text)
                st.success(f"✓ pdfminer: Extracted {debug_info['total_words']} words")
                return all_text.strip(), debug_info

        except Exception as e:
            debug_info["methods_success"]["pdfminer"] = False
            debug_info["pdfminer_error"] = str(e)
            st.warning(f"pdfminer failed: {str(e)[:100]}")

        st.error("❌ All PDF extraction methods failed")
        return "", debug_info

    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract ALL text from DOCX files - COMPLETE EXTRACTION"""
        debug_info = {"success": False}
        try:
            doc = docx.Document(file_path)

            # Extract from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                # FIXED: Include even empty paragraphs to maintain structure
                paragraphs.append(para.text if para.text else "")

            # Extract from tables
            tables_text = []
            for table_idx, table in enumerate(doc.tables):
                table_rows = []
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text if cell.text else "")
                    table_rows.append(" | ".join(row_text))

                if table_rows:
                    tables_text.append(f"\n=== Table {table_idx + 1} ===\n" + "\n".join(table_rows))

            # Extract from headers/footers
            sections_text = []
            try:
                for section_idx, section in enumerate(doc.sections):
                    if section.header:
                        for para in section.header.paragraphs:
                            if para.text.strip():
                                sections_text.append(f"Header {section_idx + 1}: {para.text}")
                    if section.footer:
                        for para in section.footer.paragraphs:
                            if para.text.strip():
                                sections_text.append(f"Footer {section_idx + 1}: {para.text}")
            except Exception as section_error:
                st.warning(f"Could not extract headers/footers: {str(section_error)[:100]}")

            # Combine all text with clear separators
            all_text_parts = []

            # Add main paragraphs
            if paragraphs:
                all_text_parts.append("\n".join(paragraphs))

            # Add tables
            if tables_text:
                all_text_parts.append("\n".join(tables_text))

            # Add sections
            if sections_text:
                all_text_parts.append("\n--- Headers & Footers ---\n" + "\n".join(sections_text))

            text = "\n\n".join(all_text_parts)

            if text.strip():
                debug_info.update({
                    "success": True,
                    "paragraphs": len(paragraphs),
                    "tables": len(doc.tables),
                    "sections": len(doc.sections),
                    "words": len(text.split()),
                    "chars": len(text)
                })
                st.success(f"✓ DOCX: {debug_info['paragraphs']} paragraphs, {debug_info['tables']} tables, {debug_info['words']} words")
                return text.strip(), debug_info
            else:
                st.error("No text found in DOCX file")
                return "", debug_info

        except Exception as e:
            st.error(f"DOCX extraction error: {str(e)[:200]}")
            debug_info["error"] = str(e)
            return "", debug_info

    @staticmethod
    def extract_text_from_txt(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract ALL text from TXT files - COMPLETE EXTRACTION with no size limits"""
        debug_info = {"success": False}

        encodings = [
            'utf-8',
            'utf-8-sig',
            'latin-1',
            'cp1252',
            'iso-8859-1',
            'ascii',
            'utf-16',
            'utf-16le',
            'utf-16be'
        ]

        file_size = os.path.getsize(file_path)
        debug_info["file_size_bytes"] = file_size
        debug_info["file_size_mb"] = round(file_size / (1024 * 1024), 2)

        # FIXED: Removed size limit, process all files
        if file_size > 10 * 1024 * 1024:  # Just a warning, not a limit
            st.info(f"📄 Large file ({debug_info['file_size_mb']} MB). Processing...")

        for encoding in encodings:
            try:
                # FIXED: Always read complete file, use larger chunks for efficiency
                text = ""
                chunk_size = 2 * 1024 * 1024  # 2MB chunks

                with open(file_path, 'r', encoding=encoding, errors='strict') as f:
                    # Read entire file in chunks to handle any size
                    chunks_read = 0
                    while True:
                        chunk = f.read(chunk_size)
                        if not chunk:
                            break
                        text += chunk
                        chunks_read += 1

                        # Progress indicator for very large files
                        if chunks_read % 10 == 0:
                            st.info(f"Processing... {chunks_read} chunks read ({chunks_read * 2} MB)")

                # Verify we got actual text
                if text.strip() and len(text.split()) > 0:
                    debug_info.update({
                        "success": True,
                        "encoding": encoding,
                        "words": len(text.split()),
                        "chars": len(text),
                        "lines": len(text.splitlines()),
                        "chunks_read": chunks_read
                    })
                    st.success(f"✓ TXT [{encoding}]: {debug_info['words']:,} words, {debug_info['lines']:,} lines ({chunks_read} chunks)")
                    return text, debug_info  # Return full text, no stripping that might lose content

            except UnicodeDecodeError:
                continue
            except Exception as e:
                st.warning(f"Encoding {encoding} failed: {str(e)[:100]}")
                continue

        # Last resort: binary read with ignore errors
        try:
            st.warning("⚠️ Using fallback method with error handling...")
            with open(file_path, 'rb') as f:
                binary_data = f.read()  # Read complete file
                text = binary_data.decode('utf-8', errors='ignore')

                if text.strip():
                    debug_info.update({
                        "success": True,
                        "encoding": "utf-8 (forced with ignore errors)",
                        "words": len(text.split()),
                        "chars": len(text),
                        "lines": len(text.splitlines())
                    })
                    st.warning(f"⚠️ TXT [fallback]: {debug_info['words']:,} words (some characters may be lost)")
                    return text, debug_info

        except Exception as e:
            st.error(f"All TXT extraction methods failed: {str(e)[:200]}")

        debug_info["error"] = "Could not decode with any encoding"
        return "", debug_info

    @classmethod
    def process_file(cls, file_path: str, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process any file type - COMPLETE EXTRACTION"""
        file_ext = filename.lower().split('.')[-1]

        # Document files
        if file_ext in ['pdf', 'txt', 'docx']:
            return cls._process_document(file_path, filename)

        # Multimedia files
        elif file_ext in ['mp4', 'avi', 'mov', 'mkv', 'flv', 'wmv',
                        'mp3', 'wav', 'ogg', 'm4a', 'flac', 'aac',
                        'jpg', 'jpeg', 'png', 'bmp', 'tiff', 'tif', 'gif']:
            return cls._process_multimedia(file_path, filename)

        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    @staticmethod
    def _process_document(file_path: str, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process document files (PDF, TXT, DOCX)"""
        file_ext = filename.lower().split('.')[-1]

        if file_ext == 'pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_ext == 'docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif file_ext == 'txt':
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            return "", {"error": f"Unsupported document type: {file_ext}"}

    @staticmethod
    def _process_multimedia(file_path: str, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process multimedia files using MultimediaProcessor"""
        return MultimediaProcessor.process_file(file_path, filename)